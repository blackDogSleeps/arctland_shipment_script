import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT


def set_font_size(step):
    if step == 0:
        f_size = Pt(30)
    elif (step == 1 or 
          step == 2 or
          step == 3):
        f_size = Pt(90)
    else:
        f_size = Pt(36)

    return f_size


def new_table(document, dictionary):
    rows_count = 8
    table = document.add_table(rows=rows_count, cols=2)
    table.style = document.styles['Table Grid']
    rows = table.rows
    step = 0

    for key, value in dictionary.items():
        rows[step].height_rule = WD_ROW_HEIGHT.AT_LEAST
        rows[step].height = Pt(70)
        f_size = set_font_size(step)
        left_cell = rows[step].cells[0]
        right_cell = rows[step].cells[1]
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        left_cell.paragraphs[0].text = key
        p = right_cell.paragraphs[0]
        p_run = p.add_run(text=str(value))
        p_run.font.size = f_size
        p_run.font.name = 'Calibri'
        p_run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # print(rows[step].cells[0].text, rows[step].cells[1].text)
        step += 1


def new_document():
    document = Document()
    font = document.styles['Normal'].font
    font.size = Pt(34)
    font.name = 'Arial'
    section = document.sections[0]
    section.left_margin = Pt(50)
    section.right_margin = Pt(35)
    section.top_margin = Pt(40)
    section.bottom_margin = Pt(1)
    return document


def get_inputs():
    boxes_total = input('Всего коробов: ')
    boxes_on_pallet = input('Коробов на палете: ')
    shipment_number = input('Номер поставки: ')
    shipment_date = input('Дата поставки: ')
    warehouse = input('Склад: ')
    pallets_amount = int(boxes_total) // int(boxes_on_pallet)
    return ({
        'Наименование юр. лица': 'ИП ХАЛЬКИН С.М. ТМ ArctLand',
        'Порядковый номер палеты': 1,
        'Кол-во палет в поставке': pallets_amount,
        'Кол-во коробок на палете': boxes_on_pallet,
        'Тип коробов': 'МИКС',
        'Номер поставки': shipment_number,
        'Дата поставки': shipment_date,
        'Склад назначения': warehouse },
        pallets_amount)


def create_output(document, shipment_date):
    output = 'output_new' 
    if output not in os.listdir():
        os.mkdir(output)
    
    document.save('output_new' + '//' + 'Маркировка ПА ' + 
                  shipment_date + '.docx')


def make_labels():
    inputs = get_inputs()
    dictionary = inputs[0]
    shipment_date = dictionary.get('Дата поставки')
    pallets_amount = inputs[1]
    document = new_document()
    pallet_index = 1
    
    for i in range(pallets_amount):
        new_table(document, dictionary)
        pallet_index += 1
        dictionary.update({'Порядковый номер палеты': pallet_index})
        if i < pallets_amount - 1:
            document.add_page_break()
    
    create_output(document, shipment_date)
    return shipment_date


def main():
    make_labels() 


if __name__ == '__main__':
    main()
